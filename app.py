from flask import Flask, send_file
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

app = Flask(__name__)

@app.route("/")
def generate_doc():

    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    # Title
    title = doc.add_paragraph("FORM ‘A’\nMEDIATION APPLICATION FORM\n")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True

    sub = doc.add_paragraph("[REFER RULE 3(1)]\nMumbai District Legal Services Authority\nCity Civil Court, Mumbai\n")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Create table with fixed column widths
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(0.6)
        row.cells[1].width = Inches(2.5)
        row.cells[2].width = Inches(3.5)

    # Header row
    hdr = table.rows[0]
    hdr.cells[0].merge(hdr.cells[2])
    p = hdr.cells[0].paragraphs[0]
    p.text = "DETAILS OF PARTIES:"
    p.runs[0].bold = True

    # Row 1
    row = table.add_row()
    row.cells[0].text = "1"
    row.cells[1].text = "Name of Applicant"
    row.cells[2].text = "{{client_name}}"

    # Row 2 (merged)
    row = table.add_row()
    row.cells[0].merge(row.cells[2])
    row.cells[0].text = "Address and contact details of Applicant"

    # Row 3
    row = table.add_row()
    row.cells[0].text = ""
    row.cells[1].text = "REGISTERED ADDRESS:"
    row.cells[2].text = "{{branch_address}}"

    # Row 4
    row = table.add_row()
    row.cells[0].text = ""
    row.cells[1].text = "CORRESPONDENCE ADDRESS:"
    row.cells[2].text = "{{branch_address}}"

    # Row 5
    row = table.add_row()
    row.cells[0].text = ""
    row.cells[1].text = "Telephone No."
    row.cells[2].text = "{{mobile}}"

    # Row 6
    row = table.add_row()
    row.cells[0].text = ""
    row.cells[1].text = "Email ID"
    row.cells[2].text = "info@kslegal.co.in"

    # Section divider
    row = table.add_row()
    row.cells[0].merge(row.cells[2])
    row.cells[0].text = "2. Name, Address and Contact details of Opposite Party"

    # ---------------- OPPOSITE PARTY DETAILS ----------------

    row = table.add_row()
    row.cells[0].merge(row.cells[2])
    row.cells[0].text = "Address and contact details of Defendant/s"

    row = table.add_row()
    row.cells[0].text = ""
    row.cells[1].text = "Name"
    row.cells[2].text = "{{customer_name}}"

    row = table.add_row()
    row.cells[0].text = ""
    row.cells[1].text = "REGISTERED ADDRESS:"
    row.cells[2].text = "{{address1}}"

    row = table.add_row()
    row.cells[0].text = ""
    row.cells[1].text = "CORRESPONDENCE ADDRESS:"
    row.cells[2].text = "{{address1}}"

    row = table.add_row()
    row.cells[0].text = ""
    row.cells[1].text = "Telephone No."
    row.cells[2].text = ""

    row = table.add_row()
    row.cells[0].text = ""
    row.cells[1].text = "Mobile No."
    row.cells[2].text = ""

    row = table.add_row()
    row.cells[0].text = ""
    row.cells[1].text = "Email ID"
    row.cells[2].text = ""

    # ---------------- DISPUTE DETAILS ----------------

    row = table.add_row()
    row.cells[0].merge(row.cells[2])
    row.cells[0].text = "DETAILS OF DISPUTE:"

    row = table.add_row()
    row.cells[0].merge(row.cells[2])
    row.cells[0].text = "THE COMM. COURTS (PRE-INSTITUTION.........SETTLEMENT) RULES, 2018"

    row = table.add_row()
    row.cells[0].merge(row.cells[2])
    row.cells[0].text = "Nature of disputes as per section 2(1)(c) of the Commercial Courts Act, 2015(4 of 2016):"


    file_path = "MAF.docx"
    doc.save(file_path)
    
    print("DOCX GENERATED SUCCESSFULLY")


    # 🔽 AUTO DOWNLOAD
    return send_file(
        file_path,
        as_attachment=True,
        download_name="MAF.docx"
    )


if __name__ == "__main__":
    app.run()
